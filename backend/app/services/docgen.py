"""Конвертация сгенерированного Markdown-документа в DOCX и PDF (кириллица/казахский)."""
import io
import re
from functools import lru_cache

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from fontTools.ttLib import TTFont
from fpdf import FPDF
from fpdf.enums import XPos, YPos

from ..config import ASSETS_DIR

# Tinos метрически совместим с Times New Roman, но его лицензия (OFL) разрешает вшивать
# шрифт в PDF. Подробности — backend/app/assets/fonts/README.md.
FONT_FAMILY = "Tinos"
FONT_REGULAR = ASSETS_DIR / "fonts" / "Tinos-Regular.ttf"
FONT_BOLD = ASSETS_DIR / "fonts" / "Tinos-Bold.ttf"

# Официальная вёрстка: кегль 12, поля по ГОСТ Р 7.0.97 (левое 20, правое 15, верх/низ 20 мм).
BODY_SIZE = 12
HEADING_SIZE = 12
TITLE_SIZE = 14
BODY_LINE_H = 6.0
HEADING_LINE_H = 6.5
TITLE_LINE_H = 7.5
MARGIN_LEFT = 20
MARGIN_TOP = 20
MARGIN_RIGHT = 15
MARGIN_BOTTOM = 20
LIST_INDENT = 6

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")

# Текст приходит от LLM и может содержать символы, которых в Tinos нет (эмодзи, стрелки,
# типографские значки). Без замены fpdf2 рисует пустой прямоугольник и пишет warning.
_PDF_CHAR_FALLBACKS = {
    "⚖": "",  # весы из дисклеймера
    "✓": "-",
    "✔": "-",
    "✅": "-",
    "✗": "-",
}


@lru_cache(maxsize=1)
def _pdf_charset() -> frozenset[str]:
    """Символы, доступные в обоих начертаниях шрифта."""
    charsets = [
        frozenset(chr(code) for code in TTFont(str(path)).getBestCmap())
        for path in (FONT_REGULAR, FONT_BOLD)
    ]
    return charsets[0] & charsets[1]


def _sanitize_for_pdf(text: str) -> str:
    """Заменяет отсутствующие в шрифте символы, остальные оставляет как есть."""
    charset = _pdf_charset()
    return "".join(char if char in charset else _PDF_CHAR_FALLBACKS.get(char, "") for char in text).strip()


def _split_bold(text: str) -> list[tuple[str, bool]]:
    """Разбивает строку на сегменты (текст, is_bold) по **...** разметке."""
    segments: list[tuple[str, bool]] = []
    pos = 0
    for match in _BOLD_RE.finditer(text):
        if match.start() > pos:
            segments.append((text[pos : match.start()], False))
        segments.append((match.group(1), True))
        pos = match.end()
    if pos < len(text):
        segments.append((text[pos:], False))
    return segments or [(text, False)]


def _parse_lines(markdown: str):
    """Итерирует (kind, text): kind ∈ title | heading | list_item | hr | paragraph."""
    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            yield "blank", ""
        elif stripped.startswith("# "):
            yield "title", stripped[2:].strip()
        elif stripped.startswith("## "):
            yield "heading", stripped[3:].strip()
        elif stripped.startswith("### "):
            yield "heading", stripped[4:].strip()
        elif stripped in {"---", "***", "___"}:
            yield "hr", ""
        elif stripped.startswith(("- ", "* ")):
            yield "list_item", stripped[2:].strip()
        else:
            yield "paragraph", stripped


def markdown_to_docx(markdown: str) -> bytes:
    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(BODY_SIZE)

    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    for kind, text in _parse_lines(markdown):
        if kind == "blank":
            continue
        if kind == "title":
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(_BOLD_RE.sub(r"\1", text).upper())
            run.bold = True
            run.font.size = Pt(TITLE_SIZE)
        elif kind == "heading":
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(_BOLD_RE.sub(r"\1", text))
            run.bold = True
            run.font.size = Pt(HEADING_SIZE)
        elif kind == "hr":
            doc.add_paragraph("_" * 40)
        else:
            paragraph = doc.add_paragraph(style="List Bullet" if kind == "list_item" else None)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for segment, is_bold in _split_bold(text):
                run = paragraph.add_run(segment)
                run.bold = is_bold

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def markdown_to_pdf(markdown: str) -> bytes:
    pdf = FPDF(format="A4")
    pdf.set_margins(MARGIN_LEFT, MARGIN_TOP, MARGIN_RIGHT)
    pdf.set_auto_page_break(auto=True, margin=MARGIN_BOTTOM)
    pdf.add_font(FONT_FAMILY, "", str(FONT_REGULAR))
    pdf.add_font(FONT_FAMILY, "B", str(FONT_BOLD))
    pdf.add_page()
    page_width = pdf.w - pdf.l_margin - pdf.r_margin

    def write(width: float, line_height: float, content: str, align: str = "L") -> None:
        # new_x/new_y обязательны: по умолчанию fpdf2 оставляет курсор справа от ячейки,
        # и следующий абзац начинается у правого поля, вылезая за пределы страницы.
        pdf.multi_cell(
            width,
            line_height,
            content,
            align=align,
            new_x=XPos.LMARGIN,
            new_y=YPos.NEXT,
        )

    for kind, text in _parse_lines(markdown):
        plain = _sanitize_for_pdf(_BOLD_RE.sub(r"\1", text))
        if kind not in {"blank", "hr"} and not plain:
            continue
        if kind == "blank":
            pdf.ln(BODY_LINE_H / 2)
        elif kind == "title":
            pdf.set_font(FONT_FAMILY, "B", TITLE_SIZE)
            write(page_width, TITLE_LINE_H, plain.upper(), align="C")
            pdf.ln(BODY_LINE_H / 2)
        elif kind == "heading":
            pdf.set_font(FONT_FAMILY, "B", HEADING_SIZE)
            write(page_width, HEADING_LINE_H, plain)
        elif kind == "hr":
            y = pdf.get_y() + 2
            pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
            pdf.ln(5)
        elif kind == "list_item":
            pdf.set_font(FONT_FAMILY, "", BODY_SIZE)
            pdf.set_x(pdf.l_margin + LIST_INDENT)
            write(page_width - LIST_INDENT, BODY_LINE_H, f"• {plain}", align="J")
        else:
            # Инлайн-жирный в выключенном по формату абзаце fpdf2 не умеет (markdown=True
            # ломает плейсхолдеры «___»), поэтому жирным рисуем строки, выделенные целиком.
            segments = _split_bold(text)
            whole_line_bold = len(segments) == 1 and segments[0][1]
            pdf.set_font(FONT_FAMILY, "B" if whole_line_bold else "", BODY_SIZE)
            write(page_width, BODY_LINE_H, plain, align="L" if whole_line_bold else "J")

    return bytes(pdf.output())
